import os
import PyPDF2
import docx
from langchain.vectorstores import FAISS
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

class RAGPipeline:
    def __init__(self):
        self.embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        self.vectorstore = None
        self.splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)

    def _load_pdf(self, file_path):
        with open(file_path, "rb") as f:
            pdf = PyPDF2.PdfReader(f)
            return "".join((page.extract_text() or "") for page in pdf.pages).strip()

    def _load_txt(self, file_path):
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read().strip()

    def _load_docx(self, file_path):
        doc = docx.Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

    def index_documents(self, file_paths):
        docs = []
        for file_path in file_paths:
            if not os.path.exists(file_path):
                print(f"⚠️ File not found: {file_path}")
                continue

            text = ""
            if file_path.endswith(".pdf"):
                text = self._load_pdf(file_path)
            elif file_path.endswith(".txt"):
                text = self._load_txt(file_path)
            elif file_path.endswith(".docx"):
                text = self._load_docx(file_path)

            if text:
                chunks = self.splitter.split_text(text)
                docs.extend([Document(page_content=chunk, metadata={"source": os.path.basename(file_path)}) for chunk in chunks])

        if docs:
            self.vectorstore = FAISS.from_documents(docs, self.embeddings)

        return self.vectorstore

    def index_transcript(self, transcript_text):
        transcript_text = transcript_text.strip()
        if not transcript_text:
            return
        chunks = self.splitter.split_text(transcript_text)
        if self.vectorstore:
            self.vectorstore.add_texts(chunks, metadatas=[{"source": "transcript"}] * len(chunks))
        else:
            self.vectorstore = FAISS.from_texts(chunks, self.embeddings, metadatas=[{"source": "transcript"}] * len(chunks))

    def retrieve(self, query, k=3):
        if self.vectorstore:
            results = self.vectorstore.similarity_search(query, k=k)
            return [{"source": doc.metadata["source"], "content": doc.page_content} for doc in results]
        return []

if __name__ == "__main__":
    rag = RAGPipeline()
    rag.index_documents(["/Users/neha/meeting-copilot/R_Jitesh_BOA.pdf"])  # You can also pass .txt or .docx
    rag.index_transcript("Meeting discussion about late delivery penalties and compliance rules.")
    results = rag.retrieve("What is the penalty for late delivery?")
    for doc in results:
        print(f"Source: {doc['source']}, Content: {doc['content'][:100]}...")
